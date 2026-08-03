"""Read-only audit of the project's existing Word and PowerPoint style signals."""

from __future__ import annotations

import json
import re
import zipfile
from collections import Counter
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "build" / "7b_audit" / "office_styles.json"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


def docx_info(path: Path) -> dict:
    doc = Document(path)
    fonts: Counter[str] = Counter()
    sizes: Counter[float] = Counter()
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name:
                fonts[run.font.name] += 1
            if run.font.size:
                sizes[round(run.font.size.pt, 1)] += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.name:
                            fonts[run.font.name] += 1
                        if run.font.size:
                            sizes[round(run.font.size.pt, 1)] += 1
    sections = []
    for section in doc.sections:
        sections.append(
            {
                "width_mm": round(section.page_width.mm, 2),
                "height_mm": round(section.page_height.mm, 2),
                "top_margin_mm": round(section.top_margin.mm, 2),
                "bottom_margin_mm": round(section.bottom_margin.mm, 2),
                "left_margin_mm": round(section.left_margin.mm, 2),
                "right_margin_mm": round(section.right_margin.mm, 2),
                "orientation": str(section.orientation),
            }
        )
    return {
        "path": str(path.relative_to(ROOT)),
        "sections": sections,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "fonts": fonts.most_common(12),
        "sizes_pt": sizes.most_common(12),
    }


def pptx_info(path: Path) -> dict:
    fonts: Counter[str] = Counter()
    sizes: Counter[float] = Counter()
    colors: Counter[str] = Counter()
    slides = 0
    masters = layouts = 0
    width = height = 0
    with zipfile.ZipFile(path) as zf:
        names = zf.namelist()
        slides = sum(bool(re.fullmatch(r"ppt/slides/slide\d+\.xml", n)) for n in names)
        masters = sum(bool(re.fullmatch(r"ppt/slideMasters/slideMaster\d+\.xml", n)) for n in names)
        layouts = sum(bool(re.fullmatch(r"ppt/slideLayouts/slideLayout\d+\.xml", n)) for n in names)
        pres = ET.fromstring(zf.read("ppt/presentation.xml"))
        ns = {"p": "http://schemas.openxmlformats.org/presentationml/2006/main"}
        sz = pres.find("p:sldSz", ns)
        if sz is not None:
            width = int(sz.attrib.get("cx", "0"))
            height = int(sz.attrib.get("cy", "0"))
        for name in names:
            if not name.endswith(".xml") or not name.startswith("ppt/"):
                continue
            try:
                root = ET.fromstring(zf.read(name))
            except ET.ParseError:
                continue
            for node in root.findall(f".//{{{A_NS}}}latin") + root.findall(f".//{{{A_NS}}}ea"):
                face = node.attrib.get("typeface")
                if face and not face.startswith("+"):
                    fonts[face] += 1
            for node in root.findall(f".//{{{A_NS}}}rPr") + root.findall(f".//{{{A_NS}}}defRPr"):
                if node.attrib.get("sz"):
                    sizes[round(int(node.attrib["sz"]) / 100, 1)] += 1
            for node in root.findall(f".//{{{A_NS}}}srgbClr"):
                value = node.attrib.get("val")
                if value:
                    colors[value.upper()] += 1
    return {
        "path": str(path.relative_to(ROOT)),
        "slides": slides,
        "width_in": round(width / 914400, 2),
        "height_in": round(height / 914400, 2),
        "masters": masters,
        "layouts": layouts,
        "fonts": fonts.most_common(12),
        "sizes_pt": sizes.most_common(15),
        "colors_rgb": colors.most_common(15),
    }


def main() -> int:
    word_files = sorted(
        {p.resolve() for p in ROOT.glob("*.docx")}
        | {p.resolve() for p in (ROOT / "references" / "mature_lessonplans").glob("*.docx")}
    )
    ppt_files = sorted((ROOT / "dist" / "ppt_revised").rglob("*.pptx"))
    payload = {
        "word": [docx_info(path) for path in word_files],
        "powerpoint": [pptx_info(path) for path in ppt_files],
    }
    OUT.parent.mkdir(parents=True, exist_ok=True)
    OUT.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[OK] word={len(word_files)} pptx={len(ppt_files)} output={OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
