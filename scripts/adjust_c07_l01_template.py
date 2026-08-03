from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = (
    ROOT
    / "lessons"
    / "第07章_相交线与平行线"
    / "7.1.1_相交线与对顶角"
    / "教学设计"
    / "7.1.1_相交线与对顶角_教学设计_模板调整版.docx"
)

CHINESE_FONT = "宋体"
HEADING_FONT = "黑体"
LATIN_FONT = "Times New Roman"


def set_run_font(run, size=10.5, bold=False, font=CHINESE_FONT, color="000000"):
    run.font.name = LATIN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), font)
    return run


def set_paragraph_spacing(paragraph, line=1.0, before=0, after=0, alignment=None):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.keep_together = True
    if alignment is not None:
        paragraph.alignment = alignment


def clear_cell(cell):
    tc = cell._tc
    for child in list(tc):
        if child.tag == qn("w:p"):
            tc.remove(child)
    tc.append(OxmlElement("w:p"))


def add_paragraph(cell, parts, size=10.0, after=1.0, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
    set_paragraph_spacing(p, line=1.0, after=after, alignment=alignment)
    for text, bold in parts:
        set_run_font(p.add_run(text), size=size, bold=bold, font=HEADING_FONT if bold else CHINESE_FONT)
    return p


def set_cell_text(cell, text, size=10.5, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, vertical=True):
    clear_cell(cell)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, line=1.0, alignment=alignment)
    set_run_font(p.add_run(text), size=size, bold=bold, font=HEADING_FONT if bold else CHINESE_FONT)
    if vertical:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, spec in edges.items():
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        for key, value in spec.items():
            node.set(qn(f"w:{key}"), str(value))


def set_table_borders(table, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")


def set_row_height(row, cm, exact=True):
    row.height = Cm(cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY if exact else WD_ROW_HEIGHT_RULE.AT_LEAST
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_widths(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid_cols = table._tbl.tblGrid.gridCol_lst
    for index, width in enumerate(widths_cm):
        table.columns[index].width = Cm(width)
        if index < len(grid_cols):
            grid_cols[index].set(qn("w:w"), str(int(Cm(width))))
        for cell in table.columns[index].cells:
            cell.width = Cm(width)


def configure_section(section, top=1.35, bottom=1.0, left=1.5, right=1.5):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)
    section.different_first_page_header_footer = False


def merge_and_text(table, row, start, end, text, **kwargs):
    cell = table.cell(row, start)
    if end > start:
        cell = cell.merge(table.cell(row, end))
    set_cell_text(cell, text, **kwargs)
    return cell


def build_page_one(doc):
    title = doc.add_paragraph()
    set_paragraph_spacing(title, line=1.0, after=21, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(title.add_run("北京市和平街第一中学课时教学设计"), size=18, bold=True, font=HEADING_FONT)

    date = doc.add_paragraph()
    set_paragraph_spacing(date, line=1.0, after=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(date.add_run("授课时间　　　　年　　月　　日　　　　　　　　　　　　第　1　页"), size=10.5)

    table = doc.add_table(rows=8, cols=9)
    set_table_widths(table, [1.52, 0.60, 0.90, 4.35, 3.54, 1.68, 1.54, 2.47, 1.40])
    set_table_borders(table, size=6)
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=85, start=100, bottom=85, end=100)

    merge_and_text(table, 0, 0, 2, "课题", bold=False)
    merge_and_text(table, 0, 3, 5, "7.1.1 两条直线相交（邻补角与对顶角）", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    merge_and_text(table, 0, 6, 6, "课型")
    merge_and_text(table, 0, 7, 8, "概念形成课", alignment=WD_ALIGN_PARAGRAPH.LEFT)

    merge_and_text(table, 1, 0, 2, "章（单元）总课时")
    merge_and_text(table, 1, 3, 3, "13")
    merge_and_text(table, 1, 4, 4, "本课题课时")
    merge_and_text(table, 1, 5, 5, "1")
    merge_and_text(table, 1, 6, 8, "本节课是第　1　课时")

    labels = ["教\n学\n目\n标", "教学\n重点", "教学\n难点", "教学\n方法", "教学\n手段", "板\n书\n设\n计"]
    contents = [
        [
            "1. 通过观察和标注相交线所成的四个角，说出邻补角、对顶角的定义，并能按位置关系正确识别。",
            "2. 借助补角的性质说明“对顶角相等”，能写出两步以内的推理依据。",
            "3. 已知相交线中一个角或相邻两角的数量关系，能计算其余角的度数并检验结果。",
            "4. 在转动木条、比较和交流中形成“先看位置、再用数量”的几何研究习惯。",
        ],
        ["识别邻补角和对顶角；运用对顶角相等解决简单角度问题。"],
        ["从角的两边关系准确说明对顶角，并把补角性质组织成规范推理。"],
        ["问题驱动；操作探究；例练结合；同伴互评。"],
        [
            "教师：两根可转动硬纸条、PPT、直尺、实物投影；学生：直尺、量角器、学案。",
            "教材：人教版《义务教育教科书·数学七年级下册》，印刷页1—3（PDF 8—10）。",
        ],
        [
            "7.1.1 相交线与对顶角",
            "1. 邻补角：公共边；另外两边互为反向延长线；和为180°。",
            "2. 对顶角：公共顶点；两边分别互为反向延长线。",
            "3. 性质：对顶角相等。推理：∠1+∠2=180°，∠3+∠2=180°，所以∠1=∠3。",
            "4. 方法：先看位置 → 选关系 → 列式 → 写依据 → 检验。",
            "5. 易错：相等的角不一定是对顶角。",
        ],
    ]
    heights = [4.2, 1.55, 1.55, 1.55, 2.05, 8.65]
    for idx, row_idx in enumerate(range(2, 8)):
        label_cell = table.cell(row_idx, 0).merge(table.cell(row_idx, 1))
        set_cell_text(label_cell, labels[idx], size=10.5)
        body = table.cell(row_idx, 2).merge(table.cell(row_idx, 8))
        clear_cell(body)
        for line_index, line in enumerate(contents[idx]):
            p = body.paragraphs[0] if line_index == 0 else body.add_paragraph()
            set_paragraph_spacing(p, line=1.05, after=2 if row_idx in (2, 7) else 0, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            set_run_font(p.add_run(line), size=10.0 if row_idx in (2, 7) else 10.5, bold=(row_idx == 7 and line_index == 0), font=HEADING_FONT if row_idx == 7 and line_index == 0 else CHINESE_FONT)
        body.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_row_height(table.rows[row_idx], heights[idx], exact=True)
    set_row_height(table.rows[0], 0.95, exact=True)
    set_row_height(table.rows[1], 0.95, exact=True)
    return table


PROCESS_PAGES = [
    [
        {
            "stage": "（一）情境引入",
            "teacher": [
                ("课标依据", "通过观察、操作和推理，认识相交线所成角的位置关系与数量关系；能用几何语言说明简单结论，在图形观察、分类和说理中发展几何直观、推理能力与应用意识。"),
                ("教材地位", "本课是七年级下册几何学习的起点，邻补角和对顶角是后续研究垂线、三线八角以及平行线判定与性质的直接基础。"),
                ("前后联系", "以前一册的角、平角、补角为前置知识，向后连接垂直、平行与初步几何证明。"),
                ("编写意图", "教材先让学生转动两根相交木条，比较四个角的位置和大小，再抽象概念并用补角性质推出对顶角相等。"),
                ("学情基础", "学生已认识角、平角、余角和补角，会进行180°内的角度计算。"),
                ("学情预判", "学生容易只凭“面对面”判断对顶角，忽略公共顶点和两边互为反向延长线；也可能把“对顶角相等”误写成“相等的角是对顶角”。"),
                ("常见错误", "有公共顶点且相等的角就是对顶角；相邻的两个角都是邻补角；求出一个角后重复用180°计算，未利用对顶角相等简化。"),
                ("问题", "展示剪刀、交叉道路和两根相交木条：这些图形中，两条直线的位置关系有什么共同点？"),
                ("组织", "明确本课研究任务：相交后形成哪些角关系。"),
                ("意图", "从生活模型聚焦“相交线所成角”，形成问题意识。"),
                ("纠错", "若学生把线段交叉与直线相交混同，补画反向延长线，强调研究对象是两条直线。"),
                ("PPT", "1—3"),
            ],
            "student": [
                ("活动", "观察并指出两条直线相交，形成四个角；用自己的语言提出想研究的问题。"),
                ("预期", "两条直线相交于一点；形成四个角；角之间可能有位置和大小关系。"),
            ],
            "time": "3分钟",
            "height": 10.20,
            "font_size": 8.15,
        },
        {
            "stage": "（二）旧知回顾",
            "teacher": [
                ("问题", "一条直线形成的平角是多少度？两个角互补是什么意思？"),
                ("组织", "让学生在图中指出反向延长线和公共边。"),
                ("说明", "平角为180°；两个角的和为180°时互补。"),
                ("意图", "激活定义邻补角和推导性质所需的旧知。"),
                ("纠错", "对“互补一定相邻”的回答给出两个不相邻但和为180°的角作反例。"),
                ("PPT", "4"),
            ],
            "student": [
                ("活动", "回答平角为180°，两个角的和为180°时互补；标注反向延长线和公共边。"),
                ("预期", "平角180°；互补是数量关系；反向延长线组成一条直线。"),
            ],
            "time": "4分钟",
            "height": 6.40,
        },
        {
            "stage": "（三）操作探究",
            "teacher": [
                ("组织", "学生转动两根硬纸条，给四个角编号1—4，观察位置并测量两组角。"),
                ("问题", "∠1与∠2有哪些共同与不同？∠1与∠3呢？转动后关系是否保持？"),
                ("说明", "比较∠1与∠2、∠1与∠3的位置和数量关系。"),
                ("意图", "让概念和性质从操作事实中自然产生。"),
                ("纠错", "对量角误差较大的小组，要求先看位置关系，再用角度和验证。"),
                ("PPT", "5"),
            ],
            "student": [
                ("活动", "动手转动、测量、记录∠1—∠4；小组比较两组角。"),
                ("预期", "∠1与∠2有公共边，另外两边互为反向延长线且和为180°；∠1与∠3有公共顶点，两边分别互为反向延长线且度数相等。"),
            ],
            "time": "8分钟",
            "height": 9.50,
        },
    ],
    [
        {
            "stage": "（四）概念形成",
            "teacher": [
                ("组织", "引导学生概括邻补角与对顶角的定义，并圈出关键词。"),
                ("问题", "依次展示正例和反例，要求按“两边关系”说明理由。"),
                ("说明", "邻补角既相邻又互补；对顶角有公共顶点，且一个角的两边分别是另一个角两边的反向延长线。"),
                ("意图", "从操作语言提升为准确的几何语言。"),
                ("纠错", "把“大小相等”从对顶角定义中删去，强调定义是位置关系。"),
                ("PPT", "6—7"),
            ],
            "student": [
                ("活动", "用“公共边、反向延长线、公共顶点”表述定义；判断图中角的关系并说明依据。"),
                ("预期", "能按位置关系准确识别邻补角和对顶角。"),
            ],
            "time": "6分钟",
            "height": 7.80,
        },
        {
            "stage": "（五）性质推导",
            "teacher": [
                ("问题", "由∠1+∠2=180°、∠3+∠2=180°，为什么能得出∠1=∠3？"),
                ("组织", "同桌互说推理链并补全依据。"),
                ("说明", "∠1和∠3都是∠2的补角，所以∠1=∠3；同理∠2=∠4。"),
                ("意图", "经历由已知事实推出新结论的初步证明过程。"),
                ("纠错", "若只写“对顶角相等”，要求补充性质的推出过程，区分定义与性质。"),
                ("PPT", "8"),
            ],
            "student": [
                ("活动", "根据同角的补角相等说明∠1=∠3；写出“因为—所以”的两步推理。"),
                ("预期", "能写出对顶角相等的规范推理。"),
            ],
            "time": "6分钟",
            "height": 7.80,
        },
        {
            "stage": "（六）例题示范",
            "teacher": [
                ("例题", "教材例1：直线a、b相交，∠1=40°，求∠2、∠3、∠4。"),
                ("问题", "先求哪个角？依据是什么？有没有更简洁的顺序？"),
                ("规范解答", "∠2=180°−40°=140°；∠3=∠1=40°；∠4=∠2=140°。"),
                ("意图", "形成“先识别关系—再列式—写依据—检验”的解题流程。"),
                ("纠错", "对把∠2算成40°的错误，让学生先在图上描出∠1与∠2的公共边和反向延长线。"),
                ("PPT", "9—11"),
            ],
            "student": [
                ("活动", "独立计算后与同桌核对，并说明每一步依据。"),
                ("预期", "先由邻补角求∠2=140°，再由对顶角相等得∠3=40°、∠4=140°。"),
            ],
            "time": "7分钟",
            "height": 9.75,
        },
    ],
    [
        {
            "stage": "（七）分层练习与检测",
            "teacher": [
                ("组织", "安排Q4—Q7，巡视记录“位置误判、和差错误、比例整体错误”三类问题；选择一份正确答案和一份典型错误投影互评。"),
                ("问题", "比例题中为什么两角和为180°？"),
                ("说明", "比例题先设两角为2x、7x，再用邻补角和为180°。"),
                ("意图", "检验目标达成，及时暴露并修正关键错误。"),
                ("纠错", "对比例题只写2+7=9的学生，追问“9份对应多少度”。"),
                ("PPT", "12—14"),
            ],
            "student": [
                ("活动", "先独立完成，再按评分点互评并订正；口头说明比例题依据。"),
                ("预期", "能正确识别、计算并解释。"),
            ],
            "time": "7分钟",
            "height": 7.35,
        },
        {
            "stage": "（八）回顾总结与作业",
            "teacher": [
                ("问题", "本节研究了什么？怎样识别？得到什么性质？哪里最容易错？"),
                ("组织", "布置基础巩固和提升思考两级作业。"),
                ("基础巩固", "完成学案Q5、Q7的订正，并用一句话写出每题依据；教材第3页练习第2题中取∠α=90°、115°分别计算其余三个角。"),
                ("提升思考", "完成学案Q8，并画图标出所得的两种角度；自拟一个“已知相交线中一个角，求其余角”的问题并给出答案。"),
                ("评价设计", "目标1：Q2、Q7共6分，达到5分视为达成；目标2：Q3、Q4中的推理依据共4分，达到3分视为达成；目标3：Q5、Q6、Q8共11分，达到9分视为达成。"),
                ("课堂观察", "操作记录完整、能在小组交流中使用“公共边、反向延长线、对顶角相等”等关键词。"),
                ("资源对应", "教学设计—PPT—学生学案—教师版共用Q1—Q8；教材例1对应Q4；当堂检测对应Q7—Q8。"),
                ("意图", "结构化知识并把方法迁移到课后任务。"),
                ("纠错", "若总结只谈计算，补问“定义依据是什么”。"),
                ("PPT", "15"),
            ],
            "student": [
                ("活动", "用关键词完成知识结构并进行自评，记录分层作业。"),
                ("预期", "先看位置识别邻补角和对顶角；邻补角和为180°，对顶角相等；相等的角不一定是对顶角。"),
            ],
            "time": "4分钟",
            "height": 8.35,
            "font_size": 8.55,
        },
    ],
]


def fill_rich_cell(cell, heading, items, size=9.5):
    clear_cell(cell)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, line=1.0, after=2, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_run_font(p.add_run(heading), size=10.0, bold=True, font=HEADING_FONT)
    for label, text in items:
        p = cell.add_paragraph()
        set_paragraph_spacing(p, line=1.0, after=1, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_run_font(p.add_run(f"{label}："), size=size, bold=True, font=HEADING_FONT)
        set_run_font(p.add_run(text), size=size, bold=False, font=CHINESE_FONT)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def build_process_page(doc, stages, final_page=False):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, top=1.15, bottom=0.9, left=1.5, right=1.5)
    rows = 1 + len(stages) + (1 if final_page else 0)
    table = doc.add_table(rows=rows, cols=4)
    set_table_widths(table, [1.50, 11.20, 3.50, 1.80])
    set_table_borders(table, size=6)
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=70, start=85, bottom=70, end=85)

    set_cell_text(table.cell(0, 0), "", size=9.5)
    set_cell_text(table.cell(0, 1), "教师教学活动设计", size=9.5)
    set_cell_text(table.cell(0, 2), "学生活动", size=9.5)
    set_cell_text(table.cell(0, 3), "估时", size=9.5)
    set_row_height(table.rows[0], 0.75, exact=True)
    set_repeat_header(table.rows[0])

    first = table.cell(1, 0)
    last = table.cell(len(stages), 0)
    process_label = first.merge(last) if len(stages) > 1 else first
    set_cell_text(process_label, "教\n学\n过\n程", size=10.5)

    for index, stage in enumerate(stages, 1):
        stage_size = stage.get("font_size", 9.25)
        fill_rich_cell(table.cell(index, 1), stage["stage"], stage["teacher"], size=stage_size)
        fill_rich_cell(table.cell(index, 2), "", stage["student"], size=max(stage_size, 8.8))
        set_cell_text(table.cell(index, 3), stage["time"], size=9.25)
        set_row_height(table.rows[index], stage.get("height", 8.20 if not final_page else 7.85), exact=True)
        # Match the C01-L01 template's uninterrupted vertical process columns.
        if index < len(stages):
            for col in (1, 2, 3):
                set_cell_border(table.cell(index, col), bottom={"val": "nil"})
                set_cell_border(table.cell(index + 1, col), top={"val": "nil"})

    if final_page:
        reflection_row = table.rows[-1]
        label = table.cell(rows - 1, 0)
        set_cell_text(label, "课\n后\n反\n思", size=10.5)
        body = table.cell(rows - 1, 1).merge(table.cell(rows - 1, 3))
        clear_cell(body)
        p = body.paragraphs[0]
        set_paragraph_spacing(p, line=1.0, after=0)
        set_row_height(reflection_row, 9.10, exact=True)
    return table


def main():
    doc = Document()
    configure_section(doc.sections[0], top=1.35, bottom=0.9, left=1.5, right=1.5)

    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    build_page_one(doc)
    build_process_page(doc, PROCESS_PAGES[0], final_page=False)
    build_process_page(doc, PROCESS_PAGES[1], final_page=False)
    build_process_page(doc, PROCESS_PAGES[2], final_page=True)

    # Remove the default empty paragraph before the first table if it is truly empty.
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
