"""Validate Grade 7B lesson plans against the approved C07-L01 DOCX contract."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
from pathlib import Path
from zipfile import ZipFile

import yaml
from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
GOLD = (
    ROOT / "lessons" / "第07章_相交线与平行线" / "7.1.1_相交线与对顶角"
    / "教学设计" / "7.1.1_相交线与对顶角_教学设计_模板调整版.docx"
)
BUILD = ROOT / "build" / "7b_gold_template_migration"


def sha(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest().upper()


def package_hashes(path: Path, names: list[str]) -> dict[str, str]:
    with ZipFile(path) as archive:
        return {name: sha(archive.read(name)) for name in names if name in archive.namelist()}


def source_records() -> list[tuple[dict, Path]]:
    records = []
    for path in (ROOT / "lessons").rglob("lesson.yml"):
        try:
            data = yaml.safe_load(path.read_text(encoding="utf-8"))
        except Exception:
            continue
        lesson_id = str((data.get("meta") or {}).get("lesson_id", ""))
        if lesson_id.startswith(("C07-", "C08-", "C09-", "C10-", "C11-", "C12-")):
            records.append((data, path))
    return sorted(records, key=lambda item: item[0]["meta"]["lesson_id"])


def text_of(doc: Document) -> str:
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def table_grid(table) -> list[int]:
    grid = table._tbl.tblGrid
    return [int(col.get(qn("w:w"))) for col in grid.gridCol_lst]


def section_signature(doc: Document) -> list[tuple[int, int, int, int, int, int]]:
    result = []
    for section in doc.sections:
        result.append(
            (
                int(section.page_width), int(section.page_height), int(section.top_margin),
                int(section.bottom_margin), int(section.left_margin), int(section.right_margin),
            )
        )
    return result


def row_heights(table) -> list[int | None]:
    return [int(row.height) if row.height is not None else None for row in table.rows]


def validate_one(data: dict, source: Path, gold_doc: Document, gold_parts: dict[str, str]) -> dict:
    lesson_id = data["meta"]["lesson_id"]
    lesson_root = source.parents[1]
    target = lesson_root / "教学设计" / f"{data['meta']['file_prefix']}_教学设计.docx"
    checks: list[tuple[str, bool, str]] = []

    def check(name: str, condition: bool, detail: str = "") -> None:
        checks.append((name, bool(condition), detail))

    check("file_exists", target.exists(), str(target))
    if not target.exists():
        return {"lesson_id": lesson_id, "path": str(target), "passed": 0, "failed": 1, "checks": checks}
    try:
        with ZipFile(target) as archive:
            bad = archive.testzip()
        check("zip_integrity", bad is None, str(bad or ""))
        doc = Document(target)
    except Exception as exc:
        check("docx_open", False, str(exc))
        return {"lesson_id": lesson_id, "path": str(target), "passed": sum(c[1] for c in checks), "failed": sum(not c[1] for c in checks), "checks": checks}

    text = text_of(doc)
    compact_text = text.replace("\n", "").replace("\r", "")
    check("sections_4", len(doc.sections) == 4, str(len(doc.sections)))
    check("tables_4", len(doc.tables) == 4, str(len(doc.tables)))
    check("section_geometry", section_signature(doc) == section_signature(gold_doc))
    if len(doc.tables) == 4:
        check("table_shapes", [(len(t.rows), len(t.columns)) for t in doc.tables] == [(8, 9), (4, 4), (4, 4), (4, 4)])
        check("table_grids", [table_grid(t) for t in doc.tables] == [table_grid(t) for t in gold_doc.tables])
        check("row_heights", [row_heights(t) for t in doc.tables] == [row_heights(t) for t in gold_doc.tables])
        check("reflection_blank", doc.tables[3].cell(3, 1).text.strip() == "")
        check("process_headers", all(value in compact_text for value in ["教师教学活动设计", "学生活动", "估时", "教学过程", "课后反思"]))

    check("school_title", "北京市和平街第一中学课时教学设计" in text)
    check("lesson_title", f"{data['meta']['section']} {data['meta']['title']}" in text)
    check("textbook_pages", f"印刷页{data['meta']['printed_pages']}（PDF {data['meta']['pdf_pages']}）" in text)
    check("objectives", all(value in text for value in data["objectives"]))
    check("key_difficulty", data["key_point"] in text and data["difficulty"] in text)
    check("analysis", all(value in text for value in [data["curriculum_basis"], data["textbook_analysis"]["position"], data["student_analysis"]["foundation"]]))
    check("stages", all(stage["stage"] in text for stage in data["flow"]))
    check("ppt_mapping", all(f"PPT：{stage['ppt']}" in text for stage in data["flow"]))
    check("time_total_45", sum(int(stage["minutes"]) for stage in data["flow"]) == 45)
    check("time_cells", all(f"{stage['minutes']}分钟" in text for stage in data["flow"]))
    check("blackboard", all(value in text for value in data["blackboard"]))
    check("no_upper_grade_leak", "正数和负数" not in text and "反比例函数" not in text)

    protected = ["word/styles.xml", "word/theme/theme1.xml", "word/fontTable.xml", "word/numbering.xml"]
    current_parts = package_hashes(target, protected)
    check("protected_template_parts", current_parts == gold_parts, json.dumps(current_parts, ensure_ascii=False))
    process_lengths = [sum(len(cell.text) for cell in table.rows[1].cells) for table in doc.tables[1:]]
    check("process_density_guard", process_lengths[0] <= 2600 and max(process_lengths[1:]) <= 1700, str(process_lengths))

    return {
        "lesson_id": lesson_id,
        "path": str(target),
        "passed": sum(c[1] for c in checks),
        "failed": sum(not c[1] for c in checks),
        "checks": checks,
        "process_lengths": process_lengths,
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--lesson-id", action="append")
    args = parser.parse_args()
    selected = set(args.lesson_id or [])
    gold_doc = Document(GOLD)
    protected = ["word/styles.xml", "word/theme/theme1.xml", "word/fontTable.xml", "word/numbering.xml"]
    gold_parts = package_hashes(GOLD, protected)
    results = []
    for data, source in source_records():
        lesson_id = data["meta"]["lesson_id"]
        if lesson_id == "C07-L01" or (selected and lesson_id not in selected):
            continue
        results.append(validate_one(data, source, gold_doc, gold_parts))

    BUILD.mkdir(parents=True, exist_ok=True)
    (BUILD / "lessonplan_validation.json").write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
    with (BUILD / "lessonplan_validation.csv").open("w", encoding="utf-8-sig", newline="") as stream:
        writer = csv.writer(stream)
        writer.writerow(["课时编号", "教学设计DOCX", "通过项", "失败项", "过程页首行字符密度"])
        for item in results:
            writer.writerow([item["lesson_id"], item["path"], item["passed"], item["failed"], item.get("process_lengths", "")])
    summary = {
        "lessons": len(results),
        "checks": sum(item["passed"] + item["failed"] for item in results),
        "passed": sum(item["passed"] for item in results),
        "failed": sum(item["failed"] for item in results),
    }
    print(json.dumps(summary, ensure_ascii=False))
    for item in results:
        if item["failed"]:
            print(f"[FAIL] {item['lesson_id']}")
            for name, ok, detail in item["checks"]:
                if not ok:
                    print(f"  - {name}: {detail}")
    return 1 if summary["failed"] else 0


if __name__ == "__main__":
    raise SystemExit(main())
