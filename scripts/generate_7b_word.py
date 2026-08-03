"""Generate editable Word teaching design and worksheet files for Grade 7B lessons.

The first implementation targets the locked C07-L01 sample and establishes the
shared Word style used by later batch generation.
"""

from __future__ import annotations

import argparse
from pathlib import Path

import yaml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
LESSON_ROOT = ROOT / "lessons" / "第07章_相交线与平行线" / "7.1.1_相交线与对顶角"
SOURCE = LESSON_ROOT / "构建文件" / "lesson.yml"

GREEN = "2F6B4F"
MINT = "E9F5EE"
BLUE = "DFF3FA"
LIGHT = "F6F9F7"
GRID = "8DA99A"
ORANGE = "F97316"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=GRID, size="6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_width(cell, width_mm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_mm * 56.7)))
    tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=10.5, bold=False, color="26352E", name="宋体") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, size=10.5, bold=False, color="26352E", align=None, after=0, before=0, line=1.05) -> None:
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.line_spacing = line
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def write_cell(cell, text: str, size=9.5, bold=False, color="26352E", fill=None, align=None) -> None:
    cell.text = ""
    if fill:
        set_cell_shading(cell, fill)
    set_cell_margins(cell)
    set_cell_border(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    parts = str(text).split("\n")
    for index, part in enumerate(parts):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        run = paragraph.add_run(part)
        set_run_font(run, size=size, bold=bold, color=color)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        if align is not None:
            paragraph.alignment = align


def set_table_borders(table) -> None:
    for row in table.rows:
        prevent_row_split(row)
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)


def add_page_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color="64748B")


def setup_document(doc: Document, title: str, kind: str) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(11)
    section.bottom_margin = Mm(12)
    section.left_margin = Mm(12)
    section.right_margin = Mm(12)
    section.header_distance = Mm(5)
    section.footer_distance = Mm(6)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.05
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "和平街第一中学  ·  七年级数学  ·  人教版七年级下册"
    format_paragraph(hp, size=9, color="64748B", align=WD_ALIGN_PARAGRAPH.CENTER)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f"{title}  ·  {kind}  ·  第 ")
    set_run_font(run, size=9, color="64748B")
    add_page_field(fp, "PAGE")
    run = fp.add_run(" 页 / 共 ")
    set_run_font(run, size=9, color="64748B")
    add_page_field(fp, "NUMPAGES")
    run = fp.add_run(" 页")
    set_run_font(run, size=9, color="64748B")


def add_title(doc: Document, text: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, size=18, bold=True, color=GREEN, name="黑体")
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(subtitle)
        set_run_font(r2, size=10, color="64748B")


def add_section_heading(doc: Document, text: str, answer=False) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    write_cell(cell, text, size=11.5, bold=True, color=GREEN if not answer else "1D4E89", fill=MINT if not answer else BLUE)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_body_paragraph(doc: Document, text: str, size=10.5, bold=False, color="26352E", indent=0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(indent)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)


def add_question(doc: Document, question: dict, teacher=False, include_diagram=None) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    set_cell_border(cell, color="C9D8D0")
    set_cell_margins(cell, top=90, bottom=90, start=110, end=110)
    p = cell.paragraphs[0]
    r = p.add_run(f"{question['id']}  {question['prompt']}")
    set_run_font(r, size=10.2, bold=True, color="26352E")
    p.paragraph_format.space_after = Pt(2)
    if include_diagram:
        dp = cell.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        marker = dp.add_run(include_diagram)
        set_run_font(marker, size=9, color="A0AEC0")
    if not teacher:
        blank_lines = 2 if question["id"] in {"Q2", "Q3", "Q4", "Q6", "Q8"} else 1
        for _ in range(blank_lines):
            bp = cell.add_paragraph("____________________________________________________________________")
            format_paragraph(bp, size=8.5, color="B6C2BC")
    else:
        ap = cell.add_paragraph()
        ar = ap.add_run(f"参考答案：{question['answer']}")
        set_run_font(ar, size=9.5, bold=True, color="1D4E89")
        ap.paragraph_format.space_after = Pt(1)
        for step in question.get("steps", []):
            sp = cell.add_paragraph()
            sr = sp.add_run(f"• {step}")
            set_run_font(sr, size=9.1, color="1D4E89")
        ep = cell.add_paragraph()
        er = ep.add_run(f"评分点：{question['score']}分。易错点：{question['error']}")
        set_run_font(er, size=8.8, color=ORANGE)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_manual_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_teaching_design(data: dict, output: Path) -> None:
    doc = Document()
    setup_document(doc, data["meta"]["title"], "教学设计")
    add_title(doc, "和平街第一中学课时教学设计", f"{data['meta']['lesson_id']}  {data['meta']['section']}")

    info = doc.add_table(rows=3, cols=4)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    labels = [
        ("课题", data["meta"]["title"], "课型", data["meta"]["lesson_type"]),
        ("教材", data["meta"]["textbook"], "课时", data["meta"]["period"]),
        ("年级", data["meta"]["grade"], "教材页码", f"印刷页 {data['meta']['printed_pages']}（PDF {data['meta']['pdf_pages']}）"),
    ]
    for r_index, row in enumerate(labels):
        for c_index, value in enumerate(row):
            is_label = c_index % 2 == 0
            write_cell(info.cell(r_index, c_index), value, size=9.3, bold=is_label, color=GREEN if is_label else "26352E", fill=MINT if is_label else "FFFFFF")
        info.rows[r_index].height = Mm(8)
        info.rows[r_index].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    set_table_borders(info)

    overview = doc.add_table(rows=7, cols=2)
    overview.alignment = WD_TABLE_ALIGNMENT.CENTER
    overview.autofit = False
    rows = [
        ("课标依据", data["curriculum_basis"]),
        ("教材分析", f"知识地位：{data['textbook_analysis']['position']}\n前后联系：{data['textbook_analysis']['connection']}\n编写意图：{data['textbook_analysis']['intent']}"),
        ("学情分析", f"已有基础：{data['student_analysis']['foundation']}\n可能困难：{data['student_analysis']['difficulties']}\n常见错误：{'；'.join(data['student_analysis']['misconceptions'])}"),
        ("教学目标", "\n".join(f"{i+1}. {value}" for i, value in enumerate(data["objectives"]))),
        ("重点与难点", f"重点：{data['key_point']}\n难点：{data['difficulty']}"),
        ("方法与准备", f"教学方法：{'、'.join(data['methods'])}\n教师准备：{'、'.join(data['preparation']['teacher'])}\n学生准备：{'、'.join(data['preparation']['student'])}"),
        ("板书设计", "\n".join(data["blackboard"])),
    ]
    for index, (label, value) in enumerate(rows):
        write_cell(overview.cell(index, 0), label, size=9.2, bold=True, color=GREEN, fill=MINT, align=WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(overview.cell(index, 1), value, size=8.6 if label in {"教材分析", "学情分析", "教学目标"} else 9.0)
        set_cell_width(overview.cell(index, 0), 26)
        set_cell_width(overview.cell(index, 1), 158)
        overview.rows[index].height = Mm([15, 27, 27, 30, 16, 19, 31][index])
        overview.rows[index].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    set_table_borders(overview)

    flow_groups = [data["flow"][0:3], data["flow"][3:6], data["flow"][6:8]]
    for page_index, stages in enumerate(flow_groups, start=2):
        add_manual_page_break(doc)
        add_title(doc, f"教学过程（第 {page_index} 页）", f"{data['meta']['title']} · 45分钟")
        process = doc.add_table(rows=1, cols=6)
        process.alignment = WD_TABLE_ALIGNMENT.CENTER
        process.autofit = False
        headers = ["教学环节", "教师活动", "学生活动/预期", "设计意图/纠错", "时间", "PPT"]
        widths = [24, 54, 48, 39, 10, 12]
        for col, header in enumerate(headers):
            write_cell(process.cell(0, col), header, size=8.8, bold=True, color="FFFFFF", fill=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_width(process.cell(0, col), widths[col])
        set_repeat_table_header(process.rows[0])
        for stage in stages:
            cells = process.add_row().cells
            teacher = "\n".join(f"• {value}" for value in stage["teacher"])
            student = "\n".join([*(f"• {value}" for value in stage["student"]), f"预期：{stage['expected']}"])
            intent = f"意图：{stage['intent']}\n纠错：{stage['correction']}"
            values = [stage["stage"], teacher, student, intent, f"{stage['minutes']}′", stage["ppt"]]
            for col, value in enumerate(values):
                write_cell(cells[col], value, size=7.7 if col in {1, 2, 3} else 8.4, bold=col == 0, color=GREEN if col == 0 else "26352E", fill=MINT if col == 0 else "FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER if col in {0, 4, 5} else None)
                set_cell_width(cells[col], widths[col])
            process.rows[-1].height = Mm(52 if page_index < 4 else 36)
            process.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        set_table_borders(process)
        if page_index == 4:
            summary = doc.add_table(rows=4, cols=2)
            summary.alignment = WD_TABLE_ALIGNMENT.CENTER
            summary.autofit = False
            content = [
                ("分层作业", f"基础巩固：{'；'.join(data['homework']['basic'])}\n提升思考：{'；'.join(data['homework']['advanced'])}"),
                ("评价设计", "\n".join(data["evaluation"])),
                ("资源对应", "教学设计—PPT—学生学案—教师版共用Q1—Q8；教材例1对应Q4；当堂检测对应Q7—Q8。"),
                ("教学反思", "\n\n\n"),
            ]
            for row_index, (label, value) in enumerate(content):
                write_cell(summary.cell(row_index, 0), label, size=8.8, bold=True, color=GREEN, fill=MINT, align=WD_ALIGN_PARAGRAPH.CENTER)
                write_cell(summary.cell(row_index, 1), value, size=7.9 if label in {"分层作业", "评价设计"} else 8.5)
                set_cell_width(summary.cell(row_index, 0), 26)
                set_cell_width(summary.cell(row_index, 1), 158)
                summary.rows[row_index].height = Mm([18, 20, 13, 37][row_index])
                summary.rows[row_index].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            set_table_borders(summary)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def add_worksheet_header(doc: Document, data: dict, teacher: bool) -> None:
    title = f"{data['meta']['file_prefix']}_{'学案教师版' if teacher else '学生学案'}"
    add_title(doc, title, f"教材印刷页 {data['meta']['printed_pages']}  ·  {data['meta'].get('period', '第1课时')}")
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    values = [
        ["班级", "____________", "姓名", "____________"],
        ["日期", "____________", "自评", "□已达成  □需巩固"],
    ]
    for row_index, row in enumerate(values):
        for col_index, value in enumerate(row):
            label = col_index % 2 == 0
            write_cell(table.cell(row_index, col_index), value, size=9.2, bold=label, color=GREEN if label else "26352E", fill=MINT if label else "FFFFFF")


def build_worksheet(data: dict, output: Path, teacher: bool) -> None:
    doc = Document()
    setup_document(doc, data["meta"]["title"], "学案教师版" if teacher else "学生学案")
    add_worksheet_header(doc, data, teacher)
    worksheet = data.get("worksheet", {})
    section_titles = worksheet.get("section_titles", {})
    topic = data["meta"]["title"]
    diagram_marker = "[[DIAGRAM_TOPIC]]" if data["meta"].get("needs_figure") else None
    add_section_heading(doc, "一、学习目标")
    for index, objective in enumerate(data["objectives"][:3], start=1):
        add_body_paragraph(doc, f"{index}. {objective}", size=9.6)
    add_section_heading(doc, "二、课前准备与旧知检测")
    add_question(doc, data["questions"][0], teacher=teacher)
    add_section_heading(doc, section_titles.get("self_study", "三、自主学习：提取关键信息"))
    add_question(doc, data["questions"][1], teacher=teacher, include_diagram=diagram_marker)
    add_section_heading(doc, section_titles.get("inquiry", "四、合作探究：形成结论"))
    add_question(doc, data["questions"][2], teacher=teacher)

    add_manual_page_break(doc)
    add_title(doc, f"{topic}学案", "第2页 · 概念形成与例题学习")
    add_section_heading(doc, "五、概念形成")
    if teacher:
        concept_lines = worksheet.get("concept_teacher", [
            "邻补角：有一条公共边，另一边互为反向延长线；两个角的和为180°。",
            "对顶角：有公共顶点，一个角的两边分别是另一个角两边的反向延长线。性质：对顶角相等。",
        ])
        for line in concept_lines:
            add_body_paragraph(doc, line, size=9.8, bold=True, color="1D4E89")
    else:
        concept_prompts = worksheet.get("concept_student", [
            "请用自己的话写出本课核心概念：________________________________________。",
            "成立条件或关键步骤：________________________________________________。",
            "本课重要结论：______________________________________________________。",
        ])
        for line in concept_prompts:
            add_body_paragraph(doc, line, size=10)
    add_body_paragraph(doc, worksheet.get("notice", "先确认条件，再选择概念、性质或方法；结论必须写清依据。"), size=9.6, bold=True, color=ORANGE)
    add_section_heading(doc, "六、例题学习：一例一练一归纳")
    add_question(doc, data["questions"][3], teacher=teacher, include_diagram=diagram_marker)
    add_section_heading(doc, "七、方法归纳")
    method = worksheet.get("method", "读题提取条件 → 选择概念或方法 → 规范求解 → 写出依据 → 回到问题检验。")
    if teacher:
        add_body_paragraph(doc, method, size=10, bold=True, color="1D4E89")
    else:
        add_body_paragraph(doc, worksheet.get("method_student", "提取________ → 选择________ → 规范求解 → 写出依据 → 检验。"), size=10)

    add_manual_page_break(doc)
    add_title(doc, f"{topic}学案", "第3页 · 分层练习与当堂检测")
    add_section_heading(doc, "八、基础练习")
    add_question(doc, data["questions"][4], teacher=teacher)
    add_section_heading(doc, "九、变式练习")
    add_question(doc, data["questions"][5], teacher=teacher, include_diagram=diagram_marker)
    add_section_heading(doc, "十、当堂检测")
    add_question(doc, data["questions"][6], teacher=teacher)
    add_question(doc, data["questions"][7], teacher=teacher)
    add_section_heading(doc, "十一、课堂小结与分层作业")
    if teacher:
        add_body_paragraph(doc, worksheet.get("summary", f"小结参考：围绕“{topic}”梳理核心概念、成立条件、规范步骤和常见错误。"), size=9.5, color="1D4E89")
    else:
        add_body_paragraph(doc, worksheet.get("summary_student", "我能用三个关键词概括本课：____________、____________、____________。"), size=9.6)
    add_body_paragraph(doc, f"基础巩固：{'；'.join(data['homework']['basic'])}", size=8.9)
    add_body_paragraph(doc, f"提升思考：{'；'.join(data['homework']['advanced'])}", size=8.9)
    add_body_paragraph(doc, "自我评价：□能识别  □能说理  □能计算  □还需订正：__________", size=9.4, bold=True, color=GREEN)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=SOURCE)
    parser.add_argument("--lesson-root", type=Path, default=LESSON_ROOT)
    parser.add_argument("--only", choices=["all", "teaching", "student", "teacher"], default="all")
    args = parser.parse_args()
    data = yaml.safe_load(args.source.read_text(encoding="utf-8"))
    prefix = data["meta"]["file_prefix"]
    teaching = args.lesson_root / "教学设计" / f"{prefix}_教学设计.docx"
    student = args.lesson_root / "学案" / f"{prefix}_学生学案.docx"
    teacher = args.lesson_root / "学案" / f"{prefix}_学案教师版.docx"
    generated = []
    if args.only in {"all", "teaching"}:
        build_teaching_design(data, teaching)
        generated.append(teaching)
    if args.only in {"all", "student"}:
        build_worksheet(data, student, teacher=False)
        generated.append(student)
    if args.only in {"all", "teacher"}:
        build_worksheet(data, teacher, teacher=True)
        generated.append(teacher)
    print("[OK] generated: " + "; ".join(str(path) for path in generated))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
