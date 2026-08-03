"""Apply the manually approved C07-L01 Word layout to Grade 7B lesson plans.

The reviewed C07-L01 DOCX is retained as the editable template.  Lesson content
comes from each lesson's locked ``构建文件/lesson.yml`` source.  The script never
changes C07-L01 and refuses to replace an existing target unless ``--overwrite``
is supplied together with a backup directory.
"""

from __future__ import annotations

import argparse
import hashlib
import shutil
from datetime import datetime
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from adjust_c07_l01_template import (
    CHINESE_FONT,
    HEADING_FONT,
    clear_cell,
    fill_rich_cell,
    set_cell_text,
    set_paragraph_spacing,
    set_run_font,
)


ROOT = Path(__file__).resolve().parents[1]
MANIFEST = ROOT / "config" / "curriculum_manifest_7b.yml"
GOLD_DOCX = (
    ROOT
    / "lessons"
    / "第07章_相交线与平行线"
    / "7.1.1_相交线与对顶角"
    / "教学设计"
    / "7.1.1_相交线与对顶角_教学设计_模板调整版.docx"
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def lesson_sources() -> list[Path]:
    result: list[tuple[str, Path]] = []
    for path in (ROOT / "lessons").rglob("lesson.yml"):
        try:
            data = yaml.safe_load(path.read_text(encoding="utf-8"))
        except Exception:
            continue
        lesson_id = str((data.get("meta") or {}).get("lesson_id", ""))
        if lesson_id.startswith(("C07-", "C08-", "C09-", "C10-", "C11-", "C12-")):
            result.append((lesson_id, path))
    return [path for _, path in sorted(result)]


def chapter_periods() -> dict[str, int]:
    manifest = yaml.safe_load(MANIFEST.read_text(encoding="utf-8"))
    return {item["chapter_id"]: int(item["proposed_periods"]) for item in manifest["chapters"]}


def clear_and_write_lines(cell, lines: list[str], *, size: float = 10.0, first_bold: bool = False) -> None:
    clear_cell(cell)
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        set_paragraph_spacing(
            paragraph,
            line=1.02,
            after=1.5 if len(lines) > 1 else 0,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
        )
        set_run_font(
            paragraph.add_run(str(line)),
            size=size,
            bold=first_bold and index == 0,
            font=HEADING_FONT if first_bold and index == 0 else CHINESE_FONT,
        )


def teacher_pairs(stage: dict, *, first_stage: bool, data: dict) -> list[tuple[str, str]]:
    pairs: list[tuple[str, str]] = []
    if first_stage:
        pairs.extend(
            [
                ("课标依据", data["curriculum_basis"]),
                ("教材地位", data["textbook_analysis"]["position"]),
                ("前后联系", data["textbook_analysis"]["connection"]),
                ("编写意图", data["textbook_analysis"]["intent"]),
                ("学情基础", data["student_analysis"]["foundation"]),
                ("学情预判", data["student_analysis"]["difficulties"]),
                ("常见错误", "；".join(data["student_analysis"]["misconceptions"])),
            ]
        )
    lines = list(stage.get("teacher") or [])
    for index, line in enumerate(lines):
        if "例题" in stage["stage"] and index == 0:
            label = "例题"
        elif "追问" in line or "问题" in line or "为什么" in line:
            label = "问题"
        elif index == 0:
            label = "组织"
        else:
            label = "说明"
        pairs.append((label, line))
    pairs.extend(
        [
            ("意图", stage["intent"]),
            ("纠错", stage["correction"]),
            ("PPT", str(stage["ppt"])),
        ]
    )
    return pairs


def fill_first_page(doc: Document, data: dict, periods: dict[str, int]) -> None:
    meta = data["meta"]
    title = f"{meta['section']} {meta['title']}"
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("授课时间"):
            paragraph.clear()
            set_paragraph_spacing(paragraph, line=1.0, after=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_run_font(
                paragraph.add_run("授课时间　　　　年　　月　　日　　　　　　　　　　　　第　1　页"),
                size=10.5,
            )
            break

    table = doc.tables[0]
    set_cell_text(table.cell(0, 3), title, size=10.5, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(0, 7), meta["lesson_type"], size=10.5, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    chapter_id = meta["lesson_id"].split("-")[0]
    set_cell_text(table.cell(1, 3), str(periods[chapter_id]), size=10.5)
    set_cell_text(table.cell(1, 5), "1", size=10.5)
    set_cell_text(table.cell(1, 6), "本节课是第　1　课时", size=10.5)

    clear_and_write_lines(table.cell(2, 2), [f"{i + 1}. {value}" for i, value in enumerate(data["objectives"])], size=9.8)
    clear_and_write_lines(table.cell(3, 2), [data["key_point"]], size=10.3)
    clear_and_write_lines(table.cell(4, 2), [data["difficulty"]], size=10.3)
    clear_and_write_lines(table.cell(5, 2), ["；".join(data["methods"]) + "。"], size=10.3)
    means = [
        "教师：" + "、".join(data["preparation"]["teacher"]) + "；学生：" + "、".join(data["preparation"]["student"]) + "。",
        f"教材：{meta['textbook']}，印刷页{meta['printed_pages']}（PDF {meta['pdf_pages']}）。",
    ]
    clear_and_write_lines(table.cell(6, 2), means, size=9.7)
    clear_and_write_lines(table.cell(7, 2), list(data["blackboard"]), size=9.7, first_bold=True)


def fill_process_pages(doc: Document, data: dict) -> None:
    groups = [data["flow"][0:3], data["flow"][3:6], data["flow"][6:8]]
    for table_index, stages in enumerate(groups, start=1):
        table = doc.tables[table_index]
        for row_index, stage in enumerate(stages, start=1):
            font_size = 8.0 if table_index == 1 and row_index == 1 else 9.0
            fill_rich_cell(
                table.cell(row_index, 1),
                f"（{['一','二','三','四','五','六','七','八'][sum(len(g) for g in groups[:table_index-1]) + row_index - 1]}）{stage['stage']}",
                teacher_pairs(stage, first_stage=(table_index == 1 and row_index == 1), data=data),
                size=font_size,
            )
            student_items = [("活动", value) for value in stage.get("student", [])]
            student_items.append(("预期", stage["expected"]))
            fill_rich_cell(
                table.cell(row_index, 2),
                "",
                student_items,
                size=max(font_size, 8.7),
            )
            set_cell_text(table.cell(row_index, 3), f"{stage['minutes']}分钟", size=9.25)

    # The retained template controls the final reflection row; keep it blank.
    reflection = doc.tables[3].cell(3, 1)
    clear_cell(reflection)


def build_one(source: Path, *, backup_root: Path, overwrite: bool) -> tuple[str, Path, Path]:
    data = yaml.safe_load(source.read_text(encoding="utf-8"))
    lesson_id = data["meta"]["lesson_id"]
    if lesson_id == "C07-L01":
        raise ValueError("C07-L01 is locked and must not be regenerated")
    lesson_root = source.parents[1]
    prefix = data["meta"]["file_prefix"]
    target = lesson_root / "教学设计" / f"{prefix}_教学设计.docx"
    pdf = target.with_suffix(".pdf")
    existing = [path for path in (target, pdf) if path.exists()]
    if existing and not overwrite:
        raise FileExistsError(f"Refusing to overwrite existing files: {existing}")
    if existing:
        lesson_backup = backup_root / lesson_id
        lesson_backup.mkdir(parents=True, exist_ok=True)
        for path in existing:
            shutil.copy2(path, lesson_backup / path.name)

    working = ROOT / "build" / "7b_gold_template_migration" / "working" / f"{lesson_id}.docx"
    working.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(GOLD_DOCX, working)
    doc = Document(working)
    fill_first_page(doc, data, chapter_periods())
    fill_process_pages(doc, data)
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    return lesson_id, target, pdf


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--lesson-id", action="append", help="Generate only the selected lesson ID; may repeat")
    parser.add_argument("--skip-lesson-id", action="append", help="Preserve the selected lesson ID; may repeat")
    parser.add_argument("--overwrite", action="store_true", help="Replace target files after backing them up")
    parser.add_argument("--backup-root", type=Path, help="Required with --overwrite")
    args = parser.parse_args()

    if not GOLD_DOCX.exists():
        raise FileNotFoundError(GOLD_DOCX)
    if args.overwrite and not args.backup_root:
        raise ValueError("--backup-root is required with --overwrite")
    backup_root = args.backup_root or (
        ROOT / "backup" / datetime.now().strftime("%Y%m%d_%H%M%S") / "7b_gold_template_migration"
    )
    selected = set(args.lesson_id or [])
    skipped_ids = set(args.skip_lesson_id or [])
    generated = 0
    for source in lesson_sources():
        data = yaml.safe_load(source.read_text(encoding="utf-8"))
        lesson_id = data["meta"]["lesson_id"]
        if lesson_id == "C07-L01" or lesson_id in skipped_ids or (selected and lesson_id not in selected):
            continue
        lid, target, pdf = build_one(source, backup_root=backup_root, overwrite=args.overwrite)
        generated += 1
        print(f"[DOCX] {lid} -> {target}")
        print(f"[PDF-PENDING] {pdf}")
    print(f"[SUMMARY] generated={generated} gold_sha256={sha256(GOLD_DOCX)} backup={backup_root}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
