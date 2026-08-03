#!/usr/bin/env python3
"""Extract mature DOCX lesson-plan structure without modifying the sources."""

from __future__ import annotations

import csv
import hashlib
import re
import shutil
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT
REFERENCE_DIR = ROOT / "references" / "mature_lessonplans"
BUILD_DIR = ROOT / "build" / "mature_lessonplan_analysis"
REPORT_DIR = ROOT / "reports"

SOURCE_NAMES = [
    "26.1.2反比例函数的图像与性质(2).docx",
    "26.2.1反比例函数的实际应用（1）.docx",
    "26.2.1反比例函数的实际应用（2）.docx",
]

SECTION_TERMS = (
    "教学目标",
    "教学重点",
    "教学难点",
    "教学过程",
    "教学环节",
    "教师活动",
    "学生活动",
    "设计意图",
    "板书设计",
    "课堂小结",
    "当堂检测",
    "作业",
    "例题",
    "练习",
)


def clean(text: str) -> str:
    text = text.replace("\u3000", " ").replace("\xa0", " ")
    return re.sub(r"\s+", " ", text).strip()


def compact_merged_row(values: list[str]) -> list[str]:
    """python-docx repeats merged-cell text; keep one logical value per run."""
    compact: list[str] = []
    for value in values:
        if compact and value == compact[-1]:
            continue
        compact.append(value)
    return compact


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def copy_sources() -> list[Path]:
    REFERENCE_DIR.mkdir(parents=True, exist_ok=True)
    copied: list[Path] = []
    for name in SOURCE_NAMES:
        src = SOURCE_DIR / name
        if not src.exists():
            raise FileNotFoundError(src)
        dst = REFERENCE_DIR / name
        if not dst.exists() or sha256(src) != sha256(dst):
            shutil.copy2(src, dst)
        if sha256(src) != sha256(dst):
            raise RuntimeError(f"Copy hash mismatch: {name}")
        copied.append(dst)
    return copied


def extract_images(docx_path: Path, out_dir: Path) -> list[Path]:
    image_dir = out_dir / "images"
    image_dir.mkdir(parents=True, exist_ok=True)
    images: list[Path] = []
    with zipfile.ZipFile(docx_path) as archive:
        for member in sorted(archive.namelist()):
            if not member.startswith("word/media/"):
                continue
            target = image_dir / Path(member).name
            target.write_bytes(archive.read(member))
            images.append(target)
    return images


def classify(text: str) -> str:
    for term in SECTION_TERMS:
        if term in text:
            return term
    if re.match(r"^(活动|问题|探究|例|练习|小结|作业)\s*[一二三四五六七八九十\d]", text):
        return "教学内容"
    return ""


def analyze(path: Path) -> dict:
    doc = Document(path)
    slug = path.stem
    out_dir = BUILD_DIR / slug
    out_dir.mkdir(parents=True, exist_ok=True)

    paragraphs: list[dict] = []
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = clean(paragraph.text)
        if not text:
            continue
        paragraphs.append(
            {
                "index": index,
                "style": paragraph.style.name if paragraph.style else "",
                "kind": classify(text),
                "text": text,
            }
        )

    tables: list[dict] = []
    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append(compact_merged_row([clean(cell.text) for cell in row.cells]))
        tables.append(
            {
                "index": table_index,
                "rows": rows,
                "row_count": len(rows),
                "column_count": max((len(row) for row in rows), default=0),
            }
        )

    images = extract_images(path, out_dir)
    all_text = "\n".join(item["text"] for item in paragraphs)
    for table in tables:
        all_text += "\n" + "\n".join(" | ".join(row) for row in table["rows"])

    counts = Counter()
    counts["例题"] = len(re.findall(r"例题|例\s*\d|例\s*[一二三四五六]", all_text))
    counts["练习"] = len(re.findall(r"练习|巩固|检测|反馈", all_text))
    counts["问题"] = len(re.findall(r"[？?]|问题", all_text))
    counts["设计意图"] = len(re.findall(r"设计意图|设计目的", all_text))
    counts["小结"] = len(re.findall(r"小结|总结|回顾", all_text))
    counts["作业"] = len(re.findall(r"作业", all_text))

    markdown: list[str] = [
        f"# {path.name}",
        "",
        f"- SHA-256：`{sha256(path)}`",
        f"- 非空正文段落：{len(paragraphs)}",
        f"- 表格：{len(tables)}",
        f"- 图片：{len(images)}",
        "",
        "## 正文",
        "",
    ]
    for item in paragraphs:
        label = f"【{item['kind']}】" if item["kind"] else ""
        markdown.append(f"{item['index']}. {label}{item['text']}")
    markdown.extend(["", "## 表格", ""])
    for table in tables:
        markdown.extend(
            [
                f"### 表格 {table['index']}（{table['row_count']}×{table['column_count']}）",
                "",
            ]
        )
        for row in table["rows"]:
            markdown.append("- " + " ｜ ".join(cell or "（空）" for cell in row))
        markdown.append("")
    markdown.extend(["## 图片", ""])
    if images:
        markdown.extend(f"- `{image.relative_to(ROOT)}`" for image in images)
    else:
        markdown.append("- 无嵌入图片")
    (out_dir / "extracted_content.md").write_text(
        "\n".join(markdown) + "\n", encoding="utf-8"
    )

    return {
        "file": path.name,
        "sha256": sha256(path),
        "paragraphs": paragraphs,
        "tables": tables,
        "images": images,
        "counts": counts,
    }


def structure_rows(result: dict) -> list[dict]:
    rows: list[dict] = []
    for item in result["paragraphs"]:
        if item["kind"]:
            rows.append(
                {
                    "source_file": result["file"],
                    "location": f"paragraph:{item['index']}",
                    "category": item["kind"],
                    "content": item["text"],
                    "adoption": "学习教学逻辑；纠正错字、重复编号和空白栏目后再采用",
                }
            )
    for table in result["tables"]:
        for row_index, row in enumerate(table["rows"], start=1):
            text = " ｜ ".join(cell for cell in row if cell)
            kind = classify(text)
            if kind or any(term in text for term in SECTION_TERMS):
                rows.append(
                    {
                        "source_file": result["file"],
                        "location": f"table:{table['index']}:row:{row_index}",
                        "category": kind or "教学结构",
                        "content": text,
                        "adoption": "学习环节组织；不复制原文缺陷",
                    }
                )
    return rows


def write_reports(results: list[dict]) -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    csv_path = REPORT_DIR / "mature_lessonplan_structure.csv"
    fieldnames = ["source_file", "location", "category", "content", "adoption"]
    with csv_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for result in results:
            writer.writerows(structure_rows(result))

    lines = [
        "# 三份成熟教案资料分析",
        "",
        "## 分析边界",
        "",
        "三份文档仅作为教师课时教学设计的教学逻辑与表达风格样本。分析保留原件，",
        "不把初三反比例函数内容直接移植到初一学案，也不继承错别字、重复编号、",
        "空白学生活动、空白估时或预写教学反思。",
        "",
        "## 来源完整性",
        "",
    ]
    for result in results:
        lines.append(
            f"- `{result['file']}`：SHA-256 `{result['sha256']}`；"
            f"{len(result['paragraphs'])} 个非空正文段落，{len(result['tables'])} 个表格，"
            f"{len(result['images'])} 幅嵌入图片。"
        )
    lines.extend(
        [
            "",
            "## 共同教学结构",
            "",
            "三份教案均以问题推进课堂，而不是把知识点并列罗列。稳定结构为：",
            "",
            "1. 复习旧知或以真实情境提出问题；",
            "2. 组织观察、列表、画图或数量分析，形成第一层结论；",
            "3. 以典型例题示范方法，紧跟同类练习；",
            "4. 通过追问、错例或变式归纳方法和适用条件；",
            "5. 用回顾性问题总结研究对象、过程与数学思想；",
            "6. 以当堂检测或分层作业完成反馈。",
            "",
            "## 可复用风格",
            "",
            "- 教师活动写具体问题与关键动作，避免只写“讲解”。",
            "- 每一例题之后安排同类型练习，再归纳方法，形成“一例一练”。",
            "- 设计意图短而直接，用于说明激活旧知、暴露错误、提炼方法或促进迁移。",
            "- 板书呈现课题、方法链、核心关系与易错提醒，而不是只有课题。",
            "- 应用题强调变量、不变量、关系、求解、检验与实际解释的完整建模链。",
            "- 课堂总结采用回顾性问题，让学生说研究内容、过程、方法和疑惑。",
            "",
            "## 不采用或主动修正",
            "",
            "- 不复制原文中的错别字、重复编号、标点与数学符号不规范。",
            "- 不保留空白的学生活动、设计意图或估时栏目。",
            "- 不把连续多例和大段完整解答强塞进四页模板。",
            "- 不预写真实课堂才可能产生的教学反思。",
            "- 不照搬初三题目难度、函数术语或应用背景到初一课程。",
            "",
            "## 三节样板的迁移原则",
            "",
            "- `C01-L03 数轴`：迁移“情境—抽象—辨析—作图—互查—归纳”的概念形成链。",
            "- `C05-L05 移项解一元一次方程`：迁移“旧知依据—尝试—错因—规则—示范—练习—检验”的技能链。",
            "- `C06-L07 角的比较与运算`：迁移“观察/操作—猜想—图形标注—说明依据—例练—归纳”的几何探究链。",
            "",
            "完整逐段、逐表抽取见 `build/mature_lessonplan_analysis/`；",
            "可筛选结构索引见 `reports/mature_lessonplan_structure.csv`。",
        ]
    )
    (REPORT_DIR / "mature_lessonplan_analysis.md").write_text(
        "\n".join(lines) + "\n", encoding="utf-8"
    )


def main() -> None:
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    paths = copy_sources()
    results = [analyze(path) for path in paths]
    write_reports(results)
    print(f"Extracted {len(results)} mature lesson plans.")


if __name__ == "__main__":
    main()
