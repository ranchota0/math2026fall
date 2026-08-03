#!/usr/bin/env python3
"""Extract ordered text, tables, images, and structure from mature DOCX lesson plans."""
from __future__ import annotations

import csv
import shutil
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from PIL import Image, ImageDraw, ImageOps


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "references" / "mature_lessonplans"
BUILD_DIR = ROOT / "build" / "mature_lessonplan_analysis"
REPORT_CSV = ROOT / "reports" / "mature_lessonplan_structure.csv"

SOURCES = [
    ("26.1.2反比例函数的图像与性质(2).docx", "26.1.2图像与性质"),
    ("26.2.1反比例函数的实际应用（1）.docx", "26.2.1实际应用1"),
    ("26.2.1反比例函数的实际应用（2）.docx", "26.2.1实际应用2"),
]


@dataclass
class StructureRow:
    document: str
    order: int
    kind: str
    style_or_shape: str
    text_preview: str
    row_count: int
    column_count: int


def iter_blocks(document: DocumentObject) -> Iterator[Paragraph | Table]:
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def normalize(text: str) -> str:
    return " ".join(text.replace("\u00a0", " ").split())


def extract_media(source: Path, target: Path) -> list[Path]:
    target.mkdir(parents=True, exist_ok=True)
    extracted: list[Path] = []
    with zipfile.ZipFile(source) as archive:
        for name in archive.namelist():
            if not name.startswith("word/media/") or name.endswith("/"):
                continue
            destination = target / Path(name).name
            with archive.open(name) as input_stream, destination.open("wb") as output_stream:
                shutil.copyfileobj(input_stream, output_stream)
            extracted.append(destination)
    return extracted


def extract_document(source: Path, slug: str) -> list[StructureRow]:
    target_dir = BUILD_DIR / slug
    target_dir.mkdir(parents=True, exist_ok=True)
    media = extract_media(source, target_dir / "media")
    document = Document(source)

    markdown = [
        f"# {source.name}",
        "",
        f"- 非空段落：{sum(1 for p in document.paragraphs if normalize(p.text))}",
        f"- 表格：{len(document.tables)}",
        f"- 内嵌图片：{len(media)}",
        "",
        "## 按文档顺序抽取",
        "",
    ]
    rows: list[StructureRow] = []
    order = 0
    for block in iter_blocks(document):
        order += 1
        if isinstance(block, Paragraph):
            text = normalize(block.text)
            if not text:
                continue
            style = block.style.name if block.style else ""
            markdown.extend([f"### 段落 {order}（{style or '无样式'}）", "", text, ""])
            rows.append(
                StructureRow(
                    document=source.name,
                    order=order,
                    kind="paragraph",
                    style_or_shape=style,
                    text_preview=text[:180],
                    row_count=0,
                    column_count=0,
                )
            )
            continue

        table_rows = []
        max_columns = 0
        for row_index, table_row in enumerate(block.rows, 1):
            cells = [normalize(cell.text) for cell in table_row.cells]
            max_columns = max(max_columns, len(cells))
            table_rows.append((row_index, cells))
        markdown.extend(
            [
                f"### 表格 {order}（{len(table_rows)} 行 × {max_columns} 列）",
                "",
            ]
        )
        for row_index, cells in table_rows:
            markdown.append(f"- 第 {row_index} 行：" + " | ".join(cells))
        markdown.append("")
        preview = " / ".join(" | ".join(cells) for _, cells in table_rows)
        rows.append(
            StructureRow(
                document=source.name,
                order=order,
                kind="table",
                style_or_shape=f"{len(table_rows)}x{max_columns}",
                text_preview=preview[:180],
                row_count=len(table_rows),
                column_count=max_columns,
            )
        )

    markdown.extend(["## 图片清单", ""])
    if media:
        markdown.extend(f"- `{path.name}`（{path.stat().st_size} bytes）" for path in media)
    else:
        markdown.append("- 无")
    markdown.append("")
    (target_dir / "extracted.md").write_text("\n".join(markdown), encoding="utf-8")
    build_contact_sheet(target_dir)
    return rows


def build_contact_sheet(target_dir: Path) -> None:
    page_images = sorted(target_dir.glob("page-*.png"))
    if not page_images:
        return
    thumbnails = []
    for page_number, path in enumerate(page_images, 1):
        image = Image.open(path).convert("RGB")
        image.thumbnail((360, 510))
        framed = ImageOps.expand(image, border=6, fill="white")
        canvas = Image.new("RGB", (framed.width, framed.height + 24), "white")
        canvas.paste(framed, (0, 0))
        ImageDraw.Draw(canvas).text((8, framed.height + 5), f"page {page_number}", fill="black")
        thumbnails.append(canvas)
    columns = min(3, len(thumbnails))
    cell_width = max(image.width for image in thumbnails)
    cell_height = max(image.height for image in thumbnails)
    rows = (len(thumbnails) + columns - 1) // columns
    sheet = Image.new("RGB", (columns * cell_width, rows * cell_height), "#d8d8d8")
    for index, thumbnail in enumerate(thumbnails):
        x = (index % columns) * cell_width
        y = (index // columns) * cell_height
        sheet.paste(thumbnail, (x, y))
    sheet.save(target_dir / "source-contact-sheet.png")


def main() -> int:
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    REPORT_CSV.parent.mkdir(parents=True, exist_ok=True)
    all_rows: list[StructureRow] = []
    for filename, slug in SOURCES:
        source = SOURCE_DIR / filename
        if not source.exists():
            raise FileNotFoundError(f"成熟教案不存在：{source}")
        all_rows.extend(extract_document(source, slug))
        print(f"[OK] extracted {filename}")

    with REPORT_CSV.open("w", encoding="utf-8-sig", newline="") as stream:
        writer = csv.DictWriter(stream, fieldnames=list(StructureRow.__annotations__))
        writer.writeheader()
        for row in all_rows:
            writer.writerow(row.__dict__)
    print(f"[OK] structure -> {REPORT_CSV.relative_to(ROOT)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
